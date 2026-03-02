"""
app.py — RAG Chatbot Web Application
Backend: Flask + LangChain + LangGraph + HuggingFace Inference API

Setup:
  pip install flask python-dotenv openai langchain langchain-community
              langchain-text-splitters langgraph pypdf python-docx
              beautifulsoup4 requests faiss-cpu sentence-transformers

.env file:
  HF_TOKEN=your_huggingface_token_here

Get your token at: https://huggingface.co/settings/tokens
(Create a fine-grained token with "Make calls to Inference Providers" permission)
"""

from __future__ import annotations
import os, math, csv, logging, threading
from pathlib import Path
from typing import List, Optional
from functools import partial

from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────────
UPLOAD_FOLDER  = Path("./uploads")
VECTOR_DB_PATH = Path("./vector_db")
ALLOWED_EXTS   = {"pdf", "docx", "txt", "md", "csv"}

# ✅ HuggingFace Inference — available models and router URL
HF_BASE_URL = "https://router.huggingface.co/v1"

AVAILABLE_MODELS = {
    "llama":   {
        "id":          "meta-llama/Llama-3.1-8B-Instruct",
        "label":       "Llama 3.1 8B",
        "description": "Meta · Fast & general purpose"
    },
    "qwen":    {
        "id":          "Qwen/Qwen2.5-72B-Instruct",
        "label":       "Qwen 2.5 72B",
        "description": "Alibaba · Large & powerful"
    },
}
DEFAULT_MODEL_KEY = "llama"

UPLOAD_FOLDER.mkdir(exist_ok=True)
VECTOR_DB_PATH.mkdir(exist_ok=True)

# ── Validate API key on startup ────────────────────────────────────────────────
_API_KEY = os.getenv("HF_TOKEN")
if not _API_KEY:
    raise EnvironmentError(
        "No HuggingFace token found. Set HF_TOKEN in your .env file.\n"
        "Get your token at: https://huggingface.co/settings/tokens"
    )

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32 MB

# ── Global JSON error handlers ─────────────────────────────────────────────────
@app.errorhandler(Exception)
def handle_exception(e):
    from werkzeug.exceptions import HTTPException
    if isinstance(e, HTTPException):
        return jsonify({"error": e.description}), e.code
    logger.error("Unhandled exception: %s", e, exc_info=True)
    return jsonify({"error": str(e)}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large. Max 32 MB."}), 413

@app.errorhandler(400)
def bad_req(e):
    return jsonify({"error": "Bad request: " + str(e)}), 400

@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Endpoint not found."}), 404

@app.errorhandler(500)
def srv_err(e):
    return jsonify({"error": "Internal server error: " + str(e)}), 500

# ── Global state ───────────────────────────────────────────────────────────────
_store: object        = None
_store_lock           = threading.Lock()
_hf_client            = None   # openai.OpenAI client pointed at HF router
_embeddings           = None   # LocalEmbeddings wrapper (sentence-transformers)
_workflow             = None
_active_model_key     = DEFAULT_MODEL_KEY   # currently selected model
_sources: List[str]   = []
_chat_history: List[dict] = []


# ══════════════════════════════════════════════════════════════════════════════
# ✅ HUGGINGFACE INFERENCE CLIENT + LLM CALL  (OpenAI-compatible router)
# ══════════════════════════════════════════════════════════════════════════════

def get_hf_client():
    """Return a cached OpenAI client pointed at the HuggingFace inference router."""
    global _hf_client
    if _hf_client is None:
        try:
            from openai import OpenAI
            _hf_client = OpenAI(
                api_key=_API_KEY,
                base_url=HF_BASE_URL,
            )
            logger.info("HuggingFace client initialised (model: %s)", HF_MODEL)
        except ImportError:
            raise RuntimeError("Missing dependency: pip install openai")
        except Exception as e:
            raise RuntimeError(f"Failed to create HuggingFace client: {e}") from e
    return _hf_client


def call_llm(prompt: str, system: str = "", history: list = None) -> str:
    """
    Call the currently selected HuggingFace model via the OpenAI-compatible router.
    - system:  system prompt
    - history: list of {"role": "user"/"assistant", "content": "..."} for conversation context
    - prompt:  the current user message
    """
    client   = get_hf_client()
    model_id = AVAILABLE_MODELS[_active_model_key]["id"]
    try:
        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        # Inject conversation history as proper role-based turns
        if history:
            for turn in history[-10:]:   # last 10 turns (5 exchanges) to stay within context
                role = turn.get("role", "user")
                if role not in ("user", "assistant"):
                    continue
                messages.append({"role": role, "content": turn.get("content", "")})
        # Current user message
        messages.append({"role": "user", "content": prompt})

        response = client.chat.completions.create(
            model=model_id,
            messages=messages,
            temperature=0.2,
            max_tokens=2048,
            stream=False,
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        logger.error("LLM call failed: %s", e)
        raise RuntimeError(f"LLM error: {e}") from e


# ══════════════════════════════════════════════════════════════════════════════
# ✅ LOCAL EMBEDDINGS  (sentence-transformers — no API key needed)
# HuggingFace router does not support embeddings via this API, so we use a fast local model.
# ══════════════════════════════════════════════════════════════════════════════

class LocalEmbeddings:
    """
    LangChain-compatible embeddings using sentence-transformers locally.
    Downloads ~90 MB model on first run, then runs fully offline.
    FAISS needs embed_documents(), embed_query(), and __call__().
    """
    EMBED_MODEL = "all-MiniLM-L6-v2"   # fast, good quality, 384-dim vectors

    def __init__(self):
        self._model = None

    def _get_model(self):
        if self._model is None:
            try:
                from sentence_transformers import SentenceTransformer
                logger.info("Loading local embedding model: %s (may download ~90MB on first run)...", self.EMBED_MODEL)
                self._model = SentenceTransformer(self.EMBED_MODEL)
                logger.info("Local embedding model loaded: %s", self.EMBED_MODEL)
            except ImportError:
                raise RuntimeError(
                    "Missing dependency: pip install sentence-transformers\n"
                    "Run: pip install sentence-transformers"
                )
            except Exception as e:
                logger.error("Failed to load embedding model: %s", e)
                raise RuntimeError(f"Embedding model load failed: {e}") from e
        return self._model

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        model = self._get_model()
        return model.encode(texts, show_progress_bar=False).tolist()

    def embed_query(self, text: str) -> List[float]:
        model = self._get_model()
        return model.encode([text], show_progress_bar=False)[0].tolist()

    def __call__(self, text: str) -> List[float]:
        """Makes the object callable — required by FAISS's embedding_function interface."""
        return self.embed_query(text)


def get_embeddings() -> LocalEmbeddings:
    global _embeddings
    if _embeddings is None:
        _embeddings = LocalEmbeddings()
        logger.info("Embeddings initialised: %s", LocalEmbeddings.EMBED_MODEL)
    return _embeddings


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_web(url: str):
    try:
        import requests as req_lib
    except ImportError:
        raise RuntimeError("Missing dependency: pip install requests")
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("Missing dependency: pip install beautifulsoup4")
    from urllib.parse import urlparse
    from langchain_core.documents import Document

    r = req_lib.get(url, timeout=15,
                    headers={"User-Agent": "Mozilla/5.0 (RAGBot/2.0)"})
    r.raise_for_status()
    soup = BeautifulSoup(r.text, "html.parser")
    for tag in soup.select("nav,footer,header,aside,script,style,.ads"):
        tag.decompose()
    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = "\n".join(
        l for l in main.get_text("\n", strip=True).splitlines() if l.strip()
    )
    if not text.strip():
        raise ValueError(f"No readable text found at {url}")
    parsed = urlparse(url)
    return [Document(page_content=text,
                     metadata={"source": url, "source_type": "web",
                               "domain": parsed.netloc, "label": parsed.netloc})]


def load_file(path: Path):
    from langchain_core.documents import Document
    ext = path.suffix.lower()

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("Missing dependency: pip install pypdf")
        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages, 1):
            txt = page.extract_text() or ""
            if txt.strip():
                docs.append(Document(page_content=txt,
                                     metadata={"source": str(path), "source_type": "pdf",
                                               "filename": path.name, "page": i,
                                               "label": path.name}))
        if not docs:
            raise ValueError(f"No extractable text in PDF: {path.name}")
        return docs

    elif ext == ".docx":
        try:
            from docx import Document as D
        except ImportError:
            raise RuntimeError("Missing dependency: pip install python-docx")
        doc = D(str(path))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text.strip():
            raise ValueError(f"No text in DOCX: {path.name}")
        return [Document(page_content=text,
                         metadata={"source": str(path), "source_type": "docx",
                                   "filename": path.name, "label": path.name})]

    elif ext in (".txt", ".md"):
        text = path.read_text(errors="replace").strip()
        if not text:
            raise ValueError(f"File is empty: {path.name}")
        stype = "markdown" if ext == ".md" else "text"
        return [Document(page_content=text,
                         metadata={"source": str(path), "source_type": stype,
                                   "filename": path.name, "label": path.name})]

    elif ext == ".csv":
        docs = []
        with open(path, errors="replace", newline="") as f:
            for i, row in enumerate(csv.DictReader(f), 2):
                parts = [f"{k}: {v}" for k, v in row.items()
                         if v and not _is_num(v)]
                if parts:
                    docs.append(Document(
                        page_content=" | ".join(parts),
                        metadata={"source": str(path), "source_type": "csv",
                                  "filename": path.name, "row": i,
                                  "label": path.name}))
        if not docs:
            raise ValueError(f"No usable text rows in CSV: {path.name}")
        return docs

    raise ValueError(f"Unsupported file extension: {ext}")


def _is_num(v):
    try: float(v); return True
    except: return False


def chunk_docs(docs):
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
    except ImportError:
        raise RuntimeError("Missing dependency: pip install langchain-text-splitters")
    return RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""],
        add_start_index=True,
    ).split_documents(docs)


def _remove_source_from_store(source: str):
    """Remove all existing chunks from a given source to prevent duplicates."""
    global _store
    if _store is None:
        return
    try:
        ids_to_delete = [
            doc_id
            for doc_id, doc in _store.docstore._dict.items()
            if doc.metadata.get("source") == source
        ]
        if ids_to_delete:
            _store.delete(ids_to_delete)
            logger.info("Removed %d old chunks for source: %s", len(ids_to_delete), source)
    except Exception as e:
        logger.warning("Could not remove old chunks for %s: %s", source, e)


def add_to_store(docs):
    global _store, _workflow
    if not docs:
        return 0
    chunks = chunk_docs(docs)
    if not chunks:
        raise ValueError("Document produced no chunks after splitting.")

    # Deduplicate: remove old chunks from same source(s) before re-adding
    sources_in_batch = {c.metadata.get("source") for c in chunks if c.metadata.get("source")}
    for src in sources_in_batch:
        _remove_source_from_store(src)

    emb = get_embeddings()
    with _store_lock:
        try:
            if _store is None:
                from langchain_community.vectorstores import FAISS
                _store = FAISS.from_documents(chunks, emb)
            else:
                _store.add_documents(chunks)
            _store.save_local(str(VECTOR_DB_PATH))
        except Exception as e:
            logger.error("Vector store error: %s", e)
            raise RuntimeError(f"Failed to index documents: {e}") from e
    _workflow = None
    logger.info("Indexed %d chunks", len(chunks))
    return len(chunks)


# ══════════════════════════════════════════════════════════════════════════════
# RETRIEVER
# ══════════════════════════════════════════════════════════════════════════════

def retrieve(query: str, top_k: int = 6):
    if _store is None:
        return [], [], "none"
    try:
        raw = _store.similarity_search_with_score(query, k=top_k)
    except Exception as e:
        logger.warning("Retrieval error: %s", e)
        return [], [], "none"
    normed = [
        (doc, float(math.exp(-s / 10)) if s > 1.0 else float(s))
        for doc, s in raw
    ]
    filtered = [(d, s) for d, s in normed if s >= 0.25]
    if not filtered:
        return [], [], "none"
    docs   = [d for d, _ in filtered]
    scores = [s for _, s in filtered]
    avg    = sum(scores) / len(scores)
    level  = "high" if avg >= 0.65 else "partial"
    return docs, scores, level


# ══════════════════════════════════════════════════════════════════════════════
# LANGGRAPH PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

from typing import TypedDict

class ChatState(TypedDict, total=False):
    query: str
    validated_query: str
    query_type: str
    retrieval_docs: list
    retrieval_scores: list
    retrieval_sources: list
    relevance_level: str
    rag_answer: str
    blended_answer: str
    fallback_answer: str
    final_answer: str
    citations: list
    error: Optional[str]
    conversation_history: list


RAG_SYS = (
    "You are a helpful assistant. Answer using the context below.\n"
    "Do NOT include any source labels, citations, or references like [filename] or [Source N] anywhere in your answer.\n"
    "Just answer naturally and directly.\n"
    "If the context does NOT contain enough information to answer the question, say:\n"
    "'My documents don't cover this — here's what I know:'\n"
    "Then give a thorough, helpful answer from your general knowledge.\n"
    "Use bullet points for multi-part answers.\n\n"
    "Context:\n{ctx}"
)
BLEND_SYS = (
    "You are a helpful assistant. You have some context that may partially answer the question.\n"
    "Do NOT include any source labels, citations, or references like [filename] or [Source N] in your answer.\n"
    "Just answer naturally and directly.\n"
    "If the context is insufficient or off-topic, say:\n"
    "'My documents don't have enough on this — here's what I know:'\n"
    "Then answer fully from your general knowledge.\n\n"
    "Context:\n{ctx}"
)
FALLBACK_SYS = (
    "You are a helpful assistant. No relevant documents were found for this question.\n"
    "Start your reply with: 'I couldn't find this in your documents, but here's what I know:'\n"
    "Then give a thorough, accurate, and helpful answer based on your general knowledge.\n"
    "Do NOT include any source labels, citations, or references in your answer.\n"
    "Be direct and informative."
)
CLASSIFY_SYS = (
    "Classify the query into exactly one word:\n"
    "'rag' → needs factual retrieval\n"
    "'conversational' → greeting or small talk\n"
    "'invalid' → empty or gibberish\n"
    "Reply with ONLY that one word."
)


def _make_label(d) -> str:
    """Build a human-readable citation label from document metadata."""
    stype    = d.metadata.get("source_type", "")
    filename = d.metadata.get("filename", "")
    page     = d.metadata.get("page", "")
    source   = d.metadata.get("source", "?")

    if stype == "web":
        # Full URL path e.g. bibek-koirala.com.np/research/publications
        try:
            from urllib.parse import urlparse
            parsed = urlparse(source)
            label  = parsed.netloc + (parsed.path.rstrip("/") or "")
            return label or source
        except Exception:
            return source
    elif filename and page:
        # e.g. thesis.pdf-p3
        return f"{filename}-p{page}"
    elif filename:
        return filename
    else:
        return str(source)[:80]


def _fmt(docs) -> str:
    if not docs:
        return "No context."
    # Pass only the raw text — no labels or filenames that the LLM might copy into its answer
    parts = [d.page_content for d in docs]
    return "\n\n---\n\n".join(parts)


def _history_text(history) -> str:
    lines = []
    for t in history[-6:]:
        role = "User" if t["role"] == "user" else "Assistant"
        lines.append(f"{role}: {t['content']}")
    return "\n".join(lines)


# ── Graph nodes ───────────────────────────────────────────────────────────────

def node_validate(state):
    q = (state.get("query") or "").strip()
    if not q:
        return {"validated_query": "", "query_type": "invalid", "error": "Empty query."}
    return {"validated_query": q[:4000], "error": None}


def node_classify(state):
    q = state.get("validated_query", "")
    if not q:
        return {"query_type": "invalid"}
    try:
        raw = call_llm(q, system=CLASSIFY_SYS).strip().lower()
        qt  = ("conversational" if "conversational" in raw
               else "invalid" if "invalid" in raw
               else "rag")
    except Exception:
        qt = "rag"
    return {"query_type": qt}


def node_retrieve(state):
    q = state.get("validated_query", "")
    docs, scores, level = retrieve(q)
    seen, sources = set(), []
    for d in docs:
        s = d.metadata.get("source", "?")
        if s not in seen:
            seen.add(s); sources.append(s)
    return {"retrieval_docs": docs, "retrieval_scores": scores,
            "retrieval_sources": sources, "relevance_level": level}


def node_check(state):
    return {"relevance_level": state.get("relevance_level", "none")}


def node_rag(state):
    ctx     = _fmt(state.get("retrieval_docs", []))
    history = state.get("conversation_history", [])
    prompt  = state.get("validated_query", "")
    try:
        return {"rag_answer": call_llm(prompt, system=RAG_SYS.format(ctx=ctx), history=history)}
    except Exception as e:
        logger.error("❌ RAG ERROR: %s", e)
        return {"rag_answer": "", "error": str(e)}


def node_blend(state):
    ctx     = _fmt(state.get("retrieval_docs", []))
    history = state.get("conversation_history", [])
    prompt  = state.get("validated_query", "")
    try:
        return {"blended_answer": call_llm(prompt, system=BLEND_SYS.format(ctx=ctx), history=history)}
    except Exception as e:
        logger.error("❌ BLEND ERROR: %s", e)
        return {"blended_answer": "", "error": str(e)}


def node_fallback(state):
    history = state.get("conversation_history", [])
    prompt  = state.get("validated_query", "")
    try:
        return {"fallback_answer": call_llm(prompt, system=FALLBACK_SYS, history=history)}
    except Exception as e:
        logger.error("❌ FALLBACK ERROR: %s", e)
        return {"fallback_answer": f"Error: {e}"}


def node_convo(state):
    history = state.get("conversation_history", [])
    prompt  = state.get("validated_query", "")
    try:
        return {"fallback_answer": call_llm(
            prompt,
            system="You are a friendly and helpful AI assistant.",
            history=history,
        )}
    except Exception as e:
        return {"fallback_answer": f"Error: {e}"}


def node_synthesize(state):
    level = state.get("relevance_level", "none")
    qtype = state.get("query_type", "rag")
    error = state.get("error")

    if qtype == "conversational":  ans = state.get("fallback_answer", "")
    elif level == "high":          ans = state.get("rag_answer", "")
    elif level == "partial":       ans = state.get("blended_answer", "")
    else:                          ans = state.get("fallback_answer", "")

    if not ans and error:
        ans = f"⚠️ AI Error: {error}"

    return {"final_answer": ans or "Unable to generate an answer."}


def node_citations(state):
    # Citations disabled — return answer unchanged with empty citations
    return {"final_answer": state.get("final_answer", ""), "citations": []}


def route_val(s): return "error" if s.get("error") else "classify"
def route_cls(s): return s.get("query_type", "rag")
def route_rel(s): return s.get("relevance_level", "none")


def build_workflow():
    try:
        from langgraph.graph import StateGraph, START, END
    except ImportError:
        raise RuntimeError("Missing dependency: pip install langgraph")

    g = StateGraph(ChatState)
    g.add_node("validate",   node_validate)
    g.add_node("classify",   node_classify)
    g.add_node("retrieve",   node_retrieve)
    g.add_node("check",      node_check)
    g.add_node("rag",        node_rag)
    g.add_node("blend",      node_blend)
    g.add_node("fallback",   node_fallback)
    g.add_node("convo",      node_convo)
    g.add_node("synthesize", node_synthesize)
    g.add_node("citations",  node_citations)

    g.add_edge(START, "validate")
    g.add_conditional_edges("validate", route_val,
                            {"error": "synthesize", "classify": "classify"})
    g.add_conditional_edges("classify", route_cls,
                            {"rag": "retrieve", "conversational": "convo",
                             "invalid": "synthesize"})
    g.add_edge("retrieve", "check")
    g.add_conditional_edges("check", route_rel,
                            {"high": "rag", "partial": "blend",
                             "low": "fallback", "none": "fallback"})
    for n in ("rag", "blend", "fallback", "convo"):
        g.add_edge(n, "synthesize")
    g.add_edge("synthesize", "citations")
    g.add_edge("citations", END)
    return g.compile()


def get_workflow():
    global _workflow
    if _workflow is None:
        _workflow = build_workflow()
    return _workflow


# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/status")
def status():
    return jsonify({
        "sources":        _sources,
        "indexed":        _store is not None,
        "chunk_count":    _get_chunk_count(),
        "model":          AVAILABLE_MODELS[_active_model_key]["id"],
        "model_key":      _active_model_key,
        "models":         AVAILABLE_MODELS,
        "embed_model":    LocalEmbeddings.EMBED_MODEL,
    })


@app.route("/api/model", methods=["POST"])
def set_model():
    """Switch the active LLM model."""
    global _active_model_key, _workflow
    data = request.get_json(silent=True) or {}
    key  = data.get("model_key", "").strip()
    if key not in AVAILABLE_MODELS:
        return jsonify({"error": f"Unknown model key '{key}'. Choose from: {list(AVAILABLE_MODELS.keys())}"}), 400
    _active_model_key = key
    _workflow = None  # rebuild workflow with new model context
    logger.info("Switched model to: %s (%s)", key, AVAILABLE_MODELS[key]["id"])
    return jsonify({"success": True, "model_key": key, "model": AVAILABLE_MODELS[key]})


def _get_chunk_count():
    try:
        return _store.index.ntotal if _store else 0
    except Exception:
        return 0


@app.route("/api/ingest/url", methods=["POST"])
def ingest_url():
    if not request.is_json:
        return jsonify({"error": "Content-Type must be application/json"}), 400
    data = request.get_json(silent=True) or {}
    url  = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    try:
        docs   = load_web(url)
        chunks = add_to_store(docs)
        label  = docs[0].metadata.get("domain", url) if docs else url
        if label not in _sources:
            _sources.append(label)
        logger.info("Ingested URL: %s (%d chunks)", label, chunks)
        return jsonify({"success": True, "chunks": chunks, "label": label})
    except Exception as e:
        logger.error("URL ingest error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/ingest/file", methods=["POST"])
def ingest_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded (field name must be 'file')"}), 400
    f = request.files["file"]
    if not f or f.filename == "":
        return jsonify({"error": "Empty file received"}), 400
    name = secure_filename(f.filename or "upload")
    ext  = Path(name).suffix.lower().lstrip(".")
    if ext not in ALLOWED_EXTS:
        return jsonify({"error": f"Unsupported type '.{ext}'. Allowed: {', '.join(sorted(ALLOWED_EXTS))}"}), 400
    save_path = UPLOAD_FOLDER / name
    try:
        f.save(str(save_path))
    except Exception as e:
        return jsonify({"error": f"Failed to save file: {e}"}), 500
    try:
        docs   = load_file(save_path)
        chunks = add_to_store(docs)
        if name not in _sources:
            _sources.append(name)
        logger.info("Ingested file: %s (%d chunks)", name, chunks)
        return jsonify({"success": True, "chunks": chunks, "label": name})
    except Exception as e:
        logger.error("File ingest error: %s", e, exc_info=True)
        save_path.unlink(missing_ok=True)
        return jsonify({"error": str(e)}), 500


@app.route("/api/chat", methods=["POST"])
def chat():
    if not request.is_json:
        return jsonify({"error": "Content-Type must be application/json"}), 400
    data  = request.get_json(silent=True) or {}
    query = (data.get("query") or "").strip()
    if not query:
        return jsonify({"error": "Empty query"}), 400

    try:
        wf = get_workflow()
    except Exception as e:
        return jsonify({"error": f"Failed to build workflow: {e}"}), 500

    init: ChatState = {
        "query":                query,
        "conversation_history": list(_chat_history),
        "validated_query":      "",
        "query_type":           "rag",
        "retrieval_docs":       [],
        "retrieval_scores":     [],
        "retrieval_sources":    [],
        "relevance_level":      "none",
        "rag_answer":           "",
        "blended_answer":       "",
        "fallback_answer":      "",
        "final_answer":         "",
        "citations":            [],
        "error":                None,
    }
    try:
        final = wf.invoke(init)
    except Exception as e:
        logger.error("Workflow error: %s", e, exc_info=True)
        return jsonify({"error": str(e)}), 500

    answer  = final.get("final_answer", "") or ""
    cites   = final.get("citations", [])
    level   = final.get("relevance_level", "none")
    qtype   = final.get("query_type", "rag")
    sources = final.get("retrieval_sources", [])

    _chat_history.append({"role": "user",     "content": query})
    _chat_history.append({"role": "assistant", "content": answer})

    return jsonify({
        "answer":     answer,
        "citations":  cites,
        "relevance":  level,
        "query_type": qtype,
        "sources":    sources,
    })


# ── Saved chat sessions (in-memory, survives within process) ─────────────────
_saved_sessions: list = []   # [{id, title, messages, saved_at}]

@app.route("/api/chat/history", methods=["GET"])
def get_history():
    return jsonify({"history": _chat_history})


@app.route("/api/chat/reset", methods=["POST"])
def reset_chat():
    """Save current chat (if non-empty) then start fresh."""
    global _chat_history
    data  = request.get_json(silent=True) or {}
    title = (data.get("title") or "").strip()

    if _chat_history:
        # Auto-generate title from first user message if none given
        if not title:
            first_user = next((m["content"] for m in _chat_history if m["role"] == "user"), "")
            title = first_user[:60] + ("…" if len(first_user) > 60 else "")

        import time, uuid as _uuid
        session = {
            "id":       str(_uuid.uuid4())[:8],
            "title":    title or "Untitled chat",
            "messages": list(_chat_history),
            "saved_at": time.strftime("%b %d, %H:%M"),
        }
        _saved_sessions.insert(0, session)   # newest first
        logger.info("Saved chat session: %s (%d messages)", session["title"], len(_chat_history))

    _chat_history.clear()
    return jsonify({"success": True, "sessions": _saved_sessions})


@app.route("/api/chat/sessions", methods=["GET"])
def get_sessions():
    return jsonify({"sessions": _saved_sessions})


@app.route("/api/chat/sessions/<session_id>", methods=["GET"])
def load_session(session_id):
    """Load a saved session back as the active chat."""
    global _chat_history
    session = next((s for s in _saved_sessions if s["id"] == session_id), None)
    if not session:
        return jsonify({"error": "Session not found"}), 404
    _chat_history = list(session["messages"])
    return jsonify({"success": True, "messages": _chat_history})


@app.route("/api/chat/sessions/<session_id>", methods=["DELETE"])
def delete_session(session_id):
    global _saved_sessions
    _saved_sessions = [s for s in _saved_sessions if s["id"] != session_id]
    return jsonify({"success": True})


@app.route("/api/sources/clear", methods=["POST"])
def clear_sources():
    global _store, _workflow
    _sources.clear()
    _chat_history.clear()
    _store    = None
    _workflow = None
    import shutil
    shutil.rmtree(str(VECTOR_DB_PATH), ignore_errors=True)
    VECTOR_DB_PATH.mkdir(exist_ok=True)
    return jsonify({"success": True})


# ── Startup: reload persisted FAISS index ─────────────────────────────────────
def _try_load_index():
    global _store, _sources
    idx = VECTOR_DB_PATH / "index.faiss"
    if idx.exists():
        try:
            from langchain_community.vectorstores import FAISS
            _store = FAISS.load_local(
                str(VECTOR_DB_PATH),
                get_embeddings(),
                allow_dangerous_deserialization=True,
            )
            logger.info("Loaded persisted FAISS index (%d vectors)", _store.index.ntotal)

            # Restore _sources from the metadata stored inside the FAISS docstore
            seen = set()
            for doc in _store.docstore._dict.values():
                label = (
                    doc.metadata.get("filename")
                    or doc.metadata.get("domain")
                    or doc.metadata.get("source", "")
                )
                if label and label not in seen:
                    seen.add(label)
                    _sources.append(label)
            logger.info("Restored %d source(s): %s", len(_sources), _sources)
        except Exception as e:
            logger.warning("Could not load index: %s", e)


if __name__ == "__main__":
    # Pre-warm embedding model so first ingest doesn't timeout
    logger.info("Pre-loading embedding model...")
    try:
        get_embeddings()._get_model()
    except Exception as e:
        logger.warning("Could not pre-load embedding model: %s", e)
    _try_load_index()
    app.run(debug=True, host="0.0.0.0", port=5000, threaded=True)
